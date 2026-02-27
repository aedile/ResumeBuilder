"""DOCX generator for resumes using python-docx.

CONSTITUTION Priority 5: Type hints, docstrings
"""

from io import BytesIO

from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Pt, RGBColor

from resume_builder.generators.constants import SUPPORTED_STYLES
from resume_builder.models.resume import Resume


class DOCXGenerator:
    """Generate DOCX resumes from Resume objects using python-docx."""

    def generate(self, resume: Resume, style: str = "classic") -> bytes:
        """Generate a DOCX resume with the specified style.

        Args:
            resume: Resume object to render.
            style: Visual style to use (classic, modern, tech, ats).
                   Affects heading font choice (serif vs sans-serif).

        Returns:
            DOCX document as raw bytes.

        Raises:
            ValueError: If style is not in SUPPORTED_STYLES.

        """
        if style not in SUPPORTED_STYLES:
            raise ValueError(
                f"Unknown style: {style}. Supported styles: {', '.join(sorted(SUPPORTED_STYLES))}"
            )

        doc = Document()
        heading_font = "Georgia" if style in {"classic", "ats"} else "Calibri"

        self._add_header(doc, resume, heading_font)
        self._add_positions(doc, resume, heading_font)
        self._add_education(doc, resume, heading_font)
        self._add_skills(doc, resume, heading_font)

        output = BytesIO()
        doc.save(output)
        return output.getvalue()

    def _add_header(self, doc: DocumentType, resume: Resume, heading_font: str) -> None:
        """Add name, headline, and optional summary to document."""
        title = doc.add_heading(resume.profile.full_name, level=1)
        for run in title.runs:
            run.font.name = heading_font
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(0, 0, 0)

        headline_para = doc.add_paragraph(resume.profile.headline)
        for run in headline_para.runs:
            run.font.size = Pt(13)
            run.italic = True

        if resume.profile.location:
            doc.add_paragraph(resume.profile.location)

        if resume.profile.summary:
            doc.add_paragraph(resume.profile.summary)

    def _add_positions(self, doc: DocumentType, resume: Resume, heading_font: str) -> None:
        """Add Professional Experience section to document."""
        if not resume.positions:
            return

        self._add_section_heading(doc, "Professional Experience", heading_font)
        for position in resume.positions:
            p = doc.add_paragraph()
            p.add_run(f"{position.title} - {position.company}").bold = True

            end = position.end_date.strftime("%B %Y") if position.end_date else "Present"
            doc.add_paragraph(f"{position.start_date.strftime('%B %Y')} - {end}")

            if position.description:
                doc.add_paragraph(position.description)

    def _add_education(self, doc: DocumentType, resume: Resume, heading_font: str) -> None:
        """Add Education section to document."""
        if not resume.education:
            return

        self._add_section_heading(doc, "Education", heading_font)
        for edu in resume.education:
            p = doc.add_paragraph()
            p.add_run(f"{edu.degree_name} - {edu.school_name}").bold = True

            if edu.start_year:
                years = (
                    f"{edu.start_year} - {edu.end_year}" if edu.end_year else str(edu.start_year)
                )
                doc.add_paragraph(years)

    def _add_skills(self, doc: DocumentType, resume: Resume, heading_font: str) -> None:
        """Add Skills section to document."""
        if not resume.skills:
            return

        self._add_section_heading(doc, "Skills", heading_font)
        doc.add_paragraph(", ".join(skill.name for skill in resume.skills))

    def _add_section_heading(self, doc: DocumentType, title: str, heading_font: str) -> None:
        """Add a level-2 section heading with the given font."""
        heading = doc.add_heading(title, level=2)
        for run in heading.runs:
            run.font.name = heading_font
