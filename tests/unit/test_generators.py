"""Tests for HTML, PDF, and DOCX generators and generator constants.

CONSTITUTION Priority 3: TDD RED Phase
CONSTITUTION Priority 4: 90%+ Coverage
CONSTITUTION Priority 9: WCAG 2.1 AA Accessibility
"""

from datetime import date
from io import BytesIO

import pytest
from docx import Document

from resume_builder.generators import SUPPORTED_STYLES
from resume_builder.generators.docx import DOCXGenerator
from resume_builder.generators.html import HTMLGenerator
from resume_builder.generators.pdf import PDFGenerator
from resume_builder.generators.protocol import GeneratorProtocol
from resume_builder.models.config import ContactInfo
from resume_builder.models.resume import Education, Position, Profile, Resume, Skill


@pytest.fixture
def sample_resume() -> Resume:
    """Create sample Resume for generator testing."""
    return Resume(
        profile=Profile(
            first_name="Alice",
            last_name="Johnson",
            headline="Product Manager",
        ),
        positions=[
            Position(
                company="StartupCo",
                title="Senior PM",
                start_date=date(2021, 3, 1),
                description="Led product development",
            )
        ],
        education=[],
        skills=[
            Skill(name="Product Management"),
            Skill(name="Agile"),
        ],
        certifications=[],
        projects=[],
        publications=[],
        languages=[],
        honors=[],
        volunteer=[],
    )


class TestHTMLGenerator:
    """Tests for HTMLGenerator class."""

    def test_generate_html_classic(self, sample_resume: Resume) -> None:
        """Generates valid HTML with classic style."""
        generator = HTMLGenerator()
        html = generator.generate(sample_resume, style="classic")

        assert html
        assert "<!DOCTYPE html>" in html or "<!doctype html>" in html
        assert "Alice Johnson" in html
        assert "Product Manager" in html

    def test_generate_html_modern(self, sample_resume: Resume) -> None:
        """Generates valid HTML with modern style."""
        html = HTMLGenerator().generate(sample_resume, style="modern")

        assert html
        assert "Alice Johnson" in html

    def test_generate_html_tech(self, sample_resume: Resume) -> None:
        """Generates valid HTML with tech style."""
        html = HTMLGenerator().generate(sample_resume, style="tech")

        assert html
        assert "Alice Johnson" in html

    def test_generate_html_ats(self, sample_resume: Resume) -> None:
        """Generates valid HTML with ATS style."""
        html = HTMLGenerator().generate(sample_resume, style="ats")

        assert html
        assert "Alice Johnson" in html

    def test_generate_html_unknown_style_raises(self, sample_resume: Resume) -> None:
        """Unknown style raises ValueError."""
        generator = HTMLGenerator()

        with pytest.raises(ValueError, match="Unknown style"):
            generator.generate(sample_resume, style="unknown")

    def test_generated_html_has_semantic_structure(self, sample_resume: Resume) -> None:
        """Generated HTML has semantic structure."""
        generator = HTMLGenerator()
        html = generator.generate(sample_resume, style="classic")

        assert "<header" in html or 'role="banner"' in html
        assert "<main" in html or 'role="main"' in html


class TestHTMLGeneratorContentRendering:
    """HTMLGenerator renders all resume content faithfully and safely."""

    def test_renders_position_company_name(self, sample_resume: Resume) -> None:
        """Position company name appears in the rendered HTML output."""
        html = HTMLGenerator().generate(sample_resume, style="classic")
        assert "StartupCo" in html

    def test_renders_position_title(self, sample_resume: Resume) -> None:
        """Position title (distinct from profile headline) appears in output."""
        html = HTMLGenerator().generate(sample_resume, style="classic")
        assert "Senior PM" in html

    def test_renders_position_date_range(self, sample_resume: Resume) -> None:
        """Position start month/year and 'Present' appear for a current position."""
        html = HTMLGenerator().generate(sample_resume, style="classic")
        assert "March 2021" in html
        assert "Present" in html

    def test_renders_skill_names(self, sample_resume: Resume) -> None:
        """All skill names appear in the rendered HTML output."""
        html = HTMLGenerator().generate(sample_resume, style="classic")
        assert "Product Management" in html
        assert "Agile" in html

    def test_renders_section_headings(self, sample_resume: Resume) -> None:
        """Experience and Skills section headings appear when data is present."""
        html = HTMLGenerator().generate(sample_resume, style="classic")
        assert "Professional Experience" in html
        assert "Skills" in html

    def test_escapes_html_in_user_content(self) -> None:
        """HTML in user-supplied fields is escaped to prevent XSS injection."""
        resume = Resume(
            profile=Profile(
                first_name="<script>alert('xss')</script>",
                last_name="Test",
                headline="Engineer",
            ),
        )
        html = HTMLGenerator().generate(resume, style="classic")

        assert "<script>" not in html
        assert "&lt;script&gt;" in html

    def test_empty_positions_section_omitted(self) -> None:
        """Professional Experience section is absent when the resume has no positions."""
        resume = Resume(
            profile=Profile(first_name="Jane", last_name="Smith", headline="Designer"),
        )
        html = HTMLGenerator().generate(resume, style="classic")
        assert "Professional Experience" not in html

    def test_empty_skills_section_omitted(self) -> None:
        """Skills section is absent when the resume has no skills."""
        resume = Resume(
            profile=Profile(first_name="Jane", last_name="Smith", headline="Designer"),
        )
        html = HTMLGenerator().generate(resume, style="classic")
        assert "Skills" not in html

    @pytest.mark.parametrize("style", sorted(SUPPORTED_STYLES))
    def test_all_styles_render_company_name(self, sample_resume: Resume, style: str) -> None:
        """Company name is present in output for every supported style."""
        html = HTMLGenerator().generate(sample_resume, style=style)
        assert "StartupCo" in html, f"Company name missing in {style!r} style"

    @pytest.mark.parametrize("style", sorted(SUPPORTED_STYLES))
    def test_all_styles_render_skills(self, sample_resume: Resume, style: str) -> None:
        """Skill names are present in output for every supported style."""
        html = HTMLGenerator().generate(sample_resume, style=style)
        assert "Product Management" in html, f"Skill 'Product Management' missing in {style!r}"
        assert "Agile" in html, f"Skill 'Agile' missing in {style!r}"

    def test_position_description_rendered(self, sample_resume: Resume) -> None:
        """Position description text appears in the rendered output."""
        html = HTMLGenerator().generate(sample_resume, style="classic")
        assert "Led product development" in html

    def test_renders_profile_location_when_present(self) -> None:
        """Profile location appears in the rendered HTML header when set."""
        resume = Resume(
            profile=Profile(
                first_name="Jane",
                last_name="Doe",
                headline="Engineer",
                location="San Francisco, CA",
            ),
        )
        html = HTMLGenerator().generate(resume, style="classic")
        assert "San Francisco, CA" in html

    def test_location_absent_when_not_set(self) -> None:
        """No location placeholder is rendered when profile.location is None."""
        resume = Resume(
            profile=Profile(first_name="Jane", last_name="Doe", headline="Engineer"),
        )
        html = HTMLGenerator().generate(resume, style="classic")
        # No stray 'None' or empty location element should appear
        assert "None" not in html


class TestGeneratorConstants:
    """Tests for the shared generator constants module."""

    def test_supported_styles_importable_from_generators_package(self) -> None:
        """SUPPORTED_STYLES is importable directly from resume_builder.generators."""
        from resume_builder.generators import SUPPORTED_STYLES  # noqa: PLC0415

        assert isinstance(SUPPORTED_STYLES, frozenset)
        assert len(SUPPORTED_STYLES) > 0

    def test_supported_styles_contains_expected_values(self) -> None:
        """SUPPORTED_STYLES contains exactly the four expected style names."""
        from resume_builder.generators import SUPPORTED_STYLES  # noqa: PLC0415

        assert {"classic", "modern", "tech", "ats"} == SUPPORTED_STYLES

    def test_supported_styles_importable_from_constants_module(self) -> None:
        """SUPPORTED_STYLES is importable from the constants submodule directly."""
        from resume_builder.generators.constants import SUPPORTED_STYLES  # noqa: PLC0415

        assert isinstance(SUPPORTED_STYLES, frozenset)

    def test_html_generator_uses_shared_supported_styles(self) -> None:
        """HTMLGenerator validates styles against the shared SUPPORTED_STYLES list."""
        from resume_builder.generators import SUPPORTED_STYLES  # noqa: PLC0415
        from resume_builder.generators.html import SUPPORTED_STYLES as HTML_STYLES  # noqa: PLC0415

        assert SUPPORTED_STYLES is HTML_STYLES


@pytest.fixture
def docx_resume() -> Resume:
    """Isolated Resume fixture for DOCX generator tests — never mutated."""
    return Resume(
        profile=Profile(
            first_name="Bob",
            last_name="Smith",
            headline="Software Engineer",
            summary="Experienced engineer",
        ),
        positions=[
            Position(
                company="TechCorp",
                title="Senior Engineer",
                start_date=date(2020, 1, 1),
                description="Built distributed systems",
            )
        ],
        education=[
            Education(
                school_name="State University",
                degree_name="Bachelor of Science",
                start_year=2012,
                end_year=2016,
            )
        ],
        skills=[
            Skill(name="Python"),
            Skill(name="Go"),
        ],
        certifications=[],
        projects=[],
        publications=[],
        languages=[],
        honors=[],
        volunteer=[],
    )


class TestPDFGenerator:
    """Tests for PDFGenerator class."""

    def test_generate_pdf_returns_bytes(self, sample_resume: Resume) -> None:
        """PDF generator returns non-empty bytes."""
        pdf_bytes = PDFGenerator().generate(sample_resume, style="classic")

        assert isinstance(pdf_bytes, bytes)
        assert len(pdf_bytes) > 0

    def test_pdf_has_valid_magic_bytes(self, sample_resume: Resume) -> None:
        """Generated bytes begin with the PDF magic number %%PDF."""
        pdf_bytes = PDFGenerator().generate(sample_resume, style="classic")
        assert pdf_bytes.startswith(b"%PDF")

    @pytest.mark.parametrize("style", sorted(SUPPORTED_STYLES))
    def test_generate_pdf_all_styles(self, sample_resume: Resume, style: str) -> None:
        """PDF generation produces valid output for every supported style."""
        pdf_bytes = PDFGenerator().generate(sample_resume, style=style)
        assert pdf_bytes.startswith(b"%PDF"), f"Invalid PDF for style {style!r}"

    def test_generate_pdf_unknown_style_raises(self, sample_resume: Resume) -> None:
        """Unknown style raises ValueError."""
        with pytest.raises(ValueError, match="Unknown style"):
            PDFGenerator().generate(sample_resume, style="unknown")


class TestDOCXGenerator:
    """Tests for DOCXGenerator class."""

    def test_generate_docx_returns_bytes(self, docx_resume: Resume) -> None:
        """DOCX generator returns non-empty bytes."""
        docx_bytes = DOCXGenerator().generate(docx_resume, style="classic")

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

    def test_docx_has_valid_magic_bytes(self, docx_resume: Resume) -> None:
        """Generated bytes begin with the ZIP/OOXML magic number PK."""
        docx_bytes = DOCXGenerator().generate(docx_resume, style="classic")
        assert docx_bytes.startswith(b"PK")

    def test_docx_contains_full_name(self, docx_resume: Resume) -> None:
        """Generated DOCX contains the resume holder's full name."""
        docx_bytes = DOCXGenerator().generate(docx_resume, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Bob Smith" in full_text

    def test_docx_contains_headline(self, docx_resume: Resume) -> None:
        """Generated DOCX contains the profile headline."""
        docx_bytes = DOCXGenerator().generate(docx_resume, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Software Engineer" in full_text

    def test_docx_contains_company_name(self, docx_resume: Resume) -> None:
        """Generated DOCX contains the employer's company name."""
        docx_bytes = DOCXGenerator().generate(docx_resume, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "TechCorp" in full_text

    def test_docx_contains_skill_names(self, docx_resume: Resume) -> None:
        """Generated DOCX contains all skill names."""
        docx_bytes = DOCXGenerator().generate(docx_resume, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Python" in full_text
        assert "Go" in full_text

    def test_docx_contains_education(self, docx_resume: Resume) -> None:
        """Generated DOCX contains education section content."""
        docx_bytes = DOCXGenerator().generate(docx_resume, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "State University" in full_text
        assert "Bachelor of Science" in full_text

    @pytest.mark.parametrize("style", sorted(SUPPORTED_STYLES))
    def test_generate_docx_all_styles(self, docx_resume: Resume, style: str) -> None:
        """DOCX generation produces valid output for every supported style."""
        docx_bytes = DOCXGenerator().generate(docx_resume, style=style)
        assert docx_bytes.startswith(b"PK"), f"Invalid DOCX for style {style!r}"

    def test_generate_docx_unknown_style_raises(self, docx_resume: Resume) -> None:
        """Unknown style raises ValueError."""
        with pytest.raises(ValueError, match="Unknown style"):
            DOCXGenerator().generate(docx_resume, style="unknown")

    def test_docx_contains_location_when_present(self) -> None:
        """Generated DOCX contains the profile location when set."""
        resume = Resume(
            profile=Profile(
                first_name="Carol",
                last_name="Reyes",
                headline="Designer",
                location="Austin, TX",
            ),
        )
        docx_bytes = DOCXGenerator().generate(resume, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Austin, TX" in full_text

    def test_docx_location_absent_when_not_set(self) -> None:
        """No stray 'None' appears in DOCX when profile.location is not set."""
        resume = Resume(
            profile=Profile(first_name="Carol", last_name="Reyes", headline="Designer"),
        )
        docx_bytes = DOCXGenerator().generate(resume, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "None" not in full_text


class TestGeneratorProtocol:
    """Both generators satisfy the GeneratorProtocol interface."""

    def test_pdf_generator_satisfies_protocol(self) -> None:
        """PDFGenerator is a structural subtype of GeneratorProtocol."""
        assert isinstance(PDFGenerator(), GeneratorProtocol)

    def test_docx_generator_satisfies_protocol(self) -> None:
        """DOCXGenerator is a structural subtype of GeneratorProtocol."""
        assert isinstance(DOCXGenerator(), GeneratorProtocol)


# ---------------------------------------------------------------------------
# P2-T11: Contact info rendering tests
# ---------------------------------------------------------------------------


@pytest.fixture
def resume_with_contact() -> Resume:
    """Resume fixture with full contact info for header rendering tests."""
    return Resume(
        profile=Profile(
            first_name="Dana",
            last_name="Rivera",
            headline="Data Scientist",
        ),
        contact_info=ContactInfo(
            email="dana.rivera@example.com",
            phone="555-0199",
            linkedin_url="https://linkedin.com/in/danarivera",
        ),
    )


@pytest.fixture
def resume_without_contact() -> Resume:
    """Resume fixture with no contact info."""
    return Resume(
        profile=Profile(
            first_name="Dana",
            last_name="Rivera",
            headline="Data Scientist",
        ),
    )


class TestHTMLContactInfoRendering:
    """HTML generator renders contact info fields in the header (P2-T11)."""

    def test_html_renders_email_when_contact_present(self, resume_with_contact: Resume) -> None:
        """Email address appears in HTML when contact_info is set."""
        html = HTMLGenerator().generate(resume_with_contact, style="classic")
        assert "dana.rivera@example.com" in html

    def test_html_email_is_mailto_link(self, resume_with_contact: Resume) -> None:
        """Email is rendered as a mailto: hyperlink."""
        html = HTMLGenerator().generate(resume_with_contact, style="classic")
        assert "mailto:dana.rivera@example.com" in html

    def test_html_renders_phone_when_present(self, resume_with_contact: Resume) -> None:
        """Phone number appears in HTML when contact_info.phone is set."""
        html = HTMLGenerator().generate(resume_with_contact, style="classic")
        assert "555-0199" in html

    def test_html_phone_is_tel_link(self, resume_with_contact: Resume) -> None:
        """Phone is rendered as a tel: hyperlink."""
        html = HTMLGenerator().generate(resume_with_contact, style="classic")
        assert "tel:555-0199" in html

    def test_html_renders_linkedin_url_when_present(self, resume_with_contact: Resume) -> None:
        """LinkedIn URL appears in HTML when contact_info.linkedin_url is set."""
        html = HTMLGenerator().generate(resume_with_contact, style="classic")
        assert "https://linkedin.com/in/danarivera" in html

    def test_html_linkedin_link_has_aria_label(self, resume_with_contact: Resume) -> None:
        """LinkedIn link has an aria-label for screen reader accessibility."""
        html = HTMLGenerator().generate(resume_with_contact, style="classic")
        assert "aria-label" in html
        assert "LinkedIn" in html

    def test_html_contact_absent_when_none(self, resume_without_contact: Resume) -> None:
        """No contact info appears when contact_info is None."""
        html = HTMLGenerator().generate(resume_without_contact, style="classic")
        assert "mailto:" not in html
        assert "tel:" not in html

    def test_html_phone_absent_when_not_set(self) -> None:
        """Phone is not rendered when contact_info.phone is None."""
        resume = Resume(
            profile=Profile(first_name="Dana", last_name="R", headline="Dev"),
            contact_info=ContactInfo(email="dana@example.com"),
        )
        html = HTMLGenerator().generate(resume, style="classic")
        assert "tel:" not in html
        assert "dana@example.com" in html

    def test_html_linkedin_absent_when_not_set(self) -> None:
        """LinkedIn link not rendered when contact_info.linkedin_url is None."""
        resume = Resume(
            profile=Profile(first_name="Dana", last_name="R", headline="Dev"),
            contact_info=ContactInfo(email="dana@example.com"),
        )
        html = HTMLGenerator().generate(resume, style="classic")
        assert "linkedin.com" not in html

    def test_html_contact_no_none_literal_without_contact(
        self, resume_without_contact: Resume
    ) -> None:
        """No stray 'None' appears in the rendered HTML when contact_info is absent."""
        html = HTMLGenerator().generate(resume_without_contact, style="classic")
        assert "None" not in html

    @pytest.mark.parametrize("style", sorted(SUPPORTED_STYLES))
    def test_all_styles_render_email(self, resume_with_contact: Resume, style: str) -> None:
        """Email renders across all supported styles."""
        html = HTMLGenerator().generate(resume_with_contact, style=style)
        assert "dana.rivera@example.com" in html, f"Email missing in {style!r}"


class TestDOCXContactInfoRendering:
    """DOCX generator renders contact info fields in the header (P2-T11)."""

    def test_docx_renders_email_when_contact_present(self, resume_with_contact: Resume) -> None:
        """Email address appears in DOCX when contact_info is set."""
        docx_bytes = DOCXGenerator().generate(resume_with_contact, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "dana.rivera@example.com" in full_text

    def test_docx_renders_phone_when_present(self, resume_with_contact: Resume) -> None:
        """Phone number appears in DOCX when contact_info.phone is set."""
        docx_bytes = DOCXGenerator().generate(resume_with_contact, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "555-0199" in full_text

    def test_docx_renders_linkedin_url_when_present(self, resume_with_contact: Resume) -> None:
        """LinkedIn URL appears in DOCX when contact_info.linkedin_url is set."""
        docx_bytes = DOCXGenerator().generate(resume_with_contact, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "linkedin.com/in/danarivera" in full_text

    def test_docx_contact_absent_when_none(self, resume_without_contact: Resume) -> None:
        """No contact info appears in DOCX when contact_info is None."""
        docx_bytes = DOCXGenerator().generate(resume_without_contact, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "@example.com" not in full_text
        assert "None" not in full_text

    def test_docx_phone_absent_when_not_set(self) -> None:
        """Phone not rendered in DOCX when contact_info.phone is None."""
        resume = Resume(
            profile=Profile(first_name="Dana", last_name="R", headline="Dev"),
            contact_info=ContactInfo(email="dana@example.com"),
        )
        docx_bytes = DOCXGenerator().generate(resume, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "555" not in full_text

    def test_docx_linkedin_absent_when_not_set(self) -> None:
        """LinkedIn URL not rendered in DOCX when contact_info.linkedin_url is None."""
        resume = Resume(
            profile=Profile(first_name="Dana", last_name="R", headline="Dev"),
            contact_info=ContactInfo(email="dana@example.com"),
        )
        docx_bytes = DOCXGenerator().generate(resume, style="classic")
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "linkedin.com" not in full_text


class TestResumeContactInfoModel:
    """Resume model accepts optional contact_info field (P2-T11)."""

    def test_resume_contact_info_defaults_to_none(self) -> None:
        """Resume.contact_info is None when not provided."""
        resume = Resume(profile=Profile(first_name="A", last_name="B", headline="C"))
        assert resume.contact_info is None

    def test_resume_accepts_contact_info(self) -> None:
        """Resume accepts a ContactInfo object."""
        resume = Resume(
            profile=Profile(first_name="A", last_name="B", headline="C"),
            contact_info=ContactInfo(email="ab@example.com"),
        )
        assert resume.contact_info is not None
        assert resume.contact_info.email == "ab@example.com"
